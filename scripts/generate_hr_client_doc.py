from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(__file__).resolve().parents[1] / "docs" / "HK_SME_HR_System_Client_Presentation_TC.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_run_font(run, size: float = 10.5, bold: bool = False, color: RGBColor | None = None) -> None:
    run.bold = bold
    run.font.name = "Microsoft JhengHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    set_run_font(run, size=10, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(document: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True)
        set_cell_shading(table.rows[0].cells[i], "D9EDE8")
        if widths:
            table.rows[0].cells[i].width = Inches(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            if widths:
                cells[i].width = Inches(widths[i])
    document.add_paragraph()


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        p = document.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(item)
        set_run_font(run)


def add_numbered(document: Document, items: list[str]) -> None:
    for item in items:
        p = document.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(item)
        set_run_font(run)


def add_section(document: Document, title: str) -> None:
    p = document.add_heading(title, level=1)
    for run in p.runs:
        set_run_font(run, size=16, bold=True, color=RGBColor(15, 118, 110))


def add_subsection(document: Document, title: str) -> None:
    p = document.add_heading(title, level=2)
    for run in p.runs:
        set_run_font(run, size=13, bold=True, color=RGBColor(30, 41, 59))


def add_para(document: Document, text: str, bold: bool = False) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, bold=bold)


def build_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = doc.styles
    styles["Normal"].font.name = "Microsoft JhengHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    styles["Normal"].font.size = Pt(10.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("香港中小企 HR 系統 MVP")
    set_run_font(run, size=26, bold=True, color=RGBColor(15, 118, 110))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("客戶簡報文件")
    set_run_font(run, size=16, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("適用對象：香港 30 人以下中小企、HR / Admin、管理層及員工自助流程")
    set_run_font(run, size=11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("文件日期：2026-04-26")
    set_run_font(run, size=10)

    add_table(
        doc,
        ["項目", "內容"],
        [
            ["目前狀態", "MVP 已可上線示範，核心 HR、請假、薪資與報表流程已串通。"],
            ["示範網址", "https://www.4mstrategy.com/hr"],
            ["系統定位", "為香港中小企提供集中式員工資料、假期、薪資、MPF、報表與稽核管理。"],
            ["技術架構", "Next.js + Tailwind CSS frontend，FastAPI + SQLModel backend，SQLite MVP database，Docker deployment。"],
        ],
        widths=[1.4, 5.4],
    )

    doc.add_page_break()

    add_section(doc, "1. 系統用途與客戶價值")
    add_para(doc, "這套 HR 系統的目標，是把香港中小企常見的人事行政、請假、薪資、MPF 及報表工作，由 Excel、WhatsApp、Email 和人工記錄，逐步集中到一個可登入、可追蹤、可匯出的系統。")
    add_para(doc, "目前版本定位為 MVP，可用作客戶示範、內部流程驗證及早期試用。它不是最終法規合規產品，但已具備足夠基礎去展示產品方向、主要工作流及未來擴展能力。")

    add_subsection(doc, "主要解決的問題")
    add_bullets(
        doc,
        [
            "員工資料分散在 Excel 或紙本，更新困難，版本不一致。",
            "請假申請、審批和結餘需要人工追蹤，容易出錯。",
            "薪資、MPF、無薪假扣款、佣金、花紅等計算缺乏清晰紀錄。",
            "IR56B、Payroll、MPF 報表需要重複整理資料。",
            "缺乏 Audit Trail，難以追蹤誰修改過員工、假期或薪資資料。",
            "員工無法自助查看個人資料、請假紀錄和薪資資料。",
        ],
    )

    add_subsection(doc, "適合的第一批客戶")
    add_bullets(
        doc,
        [
            "香港本地 5 至 30 人公司。",
            "目前用 Excel 管理員工、薪資及請假的公司。",
            "未需要大型企業 HRMS，但希望開始數碼化 HR 流程的中小企。",
            "需要繁體中文介面、香港 MPF / IR56B / 假期情境的公司。",
        ],
    )

    add_section(doc, "2. 目前已完成的功能")
    add_table(
        doc,
        ["模組", "目前完成內容", "客戶展示價值"],
        [
            ["登入與角色", "JWT login；Admin / HR / Employee 角色基礎；HR-only 導覽控制。", "展示不同身份登入後可看到不同功能，建立權限管理基礎。"],
            ["員工資料管理", "新增員工、員工列表、基本僱傭資料、薪資資料、銀行資料、工作地點、重設員工臨時密碼。", "HR 可集中建立和查看員工資料，減少 Excel 分散管理。"],
            ["公司設定", "部門、職位、工作地點、合約類型、銀行、假期類型、收入項目類型、扣款項目類型可由 HR/Admin 維護。", "避免把客戶公司資料硬編碼，方便每間公司自行調整。"],
            ["請假系統", "提交請假、HR/Admin 審批或拒絕、年假/病假/無薪假/其他、半日假、工作天及公眾假期排除。", "展示員工申請到 HR 審批的完整流程。"],
            ["假期規則", "可設定星期六是否工作天；可維護公眾假期；請假天數會計算 calendar days、public holidays excluded、actual leave days。", "支援香港公司常見工作日規則，為更完整假期政策打底。"],
            ["薪資與 MPF", "每月 payroll generation、基本薪資、津貼、gross income、taxable/non-taxable income、MPF、net salary。", "展示由員工薪資資料產生薪資結果及 MPF 計算。"],
            ["收入與扣款項目", "佣金、花紅、報銷、其他收入；缺勤、遲到、其他扣款；可標示 taxable / non-taxable、是否納入 MPF。", "支援中小企常見非固定薪資項目。"],
            ["無薪假扣款", "已批准無薪假會影響 payroll deduction，支援半日無薪假扣款。", "展示請假系統與薪資系統連動。"],
            ["離職結算", "Final pay 第一版：未發薪金、未放年假折現、代通知金、離職結算淨額。", "展示離職流程中最重要的金額計算基礎。"],
            ["薪資明細與糧單", "薪資表可點擊查看詳細計算；可下載 payslip HTML。", "客戶可看到每筆薪資如何計算，提升透明度。"],
            ["IR56B 報表", "CSV / Excel 匯出 MVP。", "展示未來可支援香港報稅資料整理。"],
            ["Audit Trail", "公司設定、員工、薪資、假期等敏感操作已開始記錄。", "提升管理層對系統可追蹤性的信心。"],
            ["UI / Mobile", "現代化後台 UI、手機版可收合 menu、公司設定 Drawer、薪資詳情 Drawer / bottom sheet。", "適合桌面和手機示範，操作體驗較接近現代 SaaS。"],
        ],
        widths=[1.25, 3.25, 2.35],
    )

    add_section(doc, "3. 已部署與可示範狀態")
    add_para(doc, "系統已部署到 Contabo VPS，並透過 Nginx 以子路徑方式提供服務。這代表目前不只是在本機開發環境運行，也已具備基本雲端展示能力。")
    add_table(
        doc,
        ["項目", "狀態"],
        [
            ["Frontend URL", "https://www.4mstrategy.com/hr"],
            ["Backend health check", "https://www.4mstrategy.com/hr-api/health"],
            ["Frontend container", "hr_frontend，127.0.0.1:5300 -> container:3000"],
            ["Backend container", "hr_backend，127.0.0.1:5301 -> container:8000"],
            ["Database", "SQLite MVP，host path: ~/apps/hr_system/data/hr_mvp.db"],
            ["Deployment method", "Docker Compose + Nginx reverse proxy"],
            ["Latest deployed commit", "6f3aee3 Add production deployment for Contabo"],
            ["Verified", "Frontend 200、static assets 200、health check 200、login API OK"],
        ],
        widths=[2.0, 4.8],
    )

    add_subsection(doc, "建議示範流程")
    add_numbered(
        doc,
        [
            "登入系統，展示 Admin / HR 工作台。",
            "進入公司設定，展示部門、職位、銀行、假期類型、收入項目和扣款項目可自行維護。",
            "新增或查看員工資料，展示員工基本資料、薪資資料、工作地點及銀行資料。",
            "提交請假申請，展示工作天、公眾假期、半日假及審批流程。",
            "新增收入或扣款項目，例如佣金、花紅、遲到扣款。",
            "生成薪資，點擊薪資表查看詳細計算、MPF、扣款、net salary。",
            "下載 payslip，展示員工薪資資料輸出能力。",
            "匯出 IR56B CSV / Excel MVP，展示未來報表方向。",
            "查看 Audit Trail，展示系統記錄敏感操作。",
        ],
    )

    add_section(doc, "4. 技術架構概要")
    add_table(
        doc,
        ["層級", "技術", "說明"],
        [
            ["Frontend", "Next.js + Tailwind CSS", "使用 App Router 建立 HR 後台介面，支援 /hr basePath 部署。"],
            ["Backend", "Python FastAPI", "提供 REST API、JWT login、HR business logic、報表下載。"],
            ["ORM", "SQLModel", "結合 SQLAlchemy 與 Pydantic，管理資料模型與資料庫操作。"],
            ["Database", "SQLite", "MVP 階段使用單檔資料庫，方便快速部署與展示。正式版建議轉 PostgreSQL。"],
            ["Authentication", "JWT", "目前 token 儲存在 localStorage。正式版建議改為 httpOnly cookie + refresh token。"],
            ["Export", "CSV / Excel / HTML Payslip", "支援 IR56B MVP 匯出與薪資單下載。"],
            ["Deployment", "Docker Compose + Nginx + SSL", "已部署到 VPS，透過 www.4mstrategy.com/hr 對外展示。"],
        ],
        widths=[1.25, 1.8, 3.75],
    )

    add_section(doc, "5. 目前限制與風險")
    add_para(doc, "目前系統是 MVP，適合用作示範、流程驗證和早期試用。若要變成正式商業產品，需要補強安全、合規、資料保留及更完整的香港 HR 規則。")
    add_bullets(
        doc,
        [
            "SQLite 適合 MVP，不適合多人高併發或正式長期 production；正式版建議轉 PostgreSQL。",
            "JWT 目前儲存在 localStorage；正式版建議使用 httpOnly cookie、refresh token、session timeout 和 rate limiting。",
            "HKID、銀行戶口、薪資等敏感資料目前未完全加密；正式版需要欄位級加密、遮罩及更嚴格權限。",
            "MPF、IR56B、病假、產假、侍產假、法定假日薪酬等仍需由香港 HR / 會計 / 法律專業人士覆核。",
            "IR56B 目前是 MVP 匯出，不等同完整 eTAX 或正式 IRD 格式。",
            "Payroll lock / approval flow 尚未完整，正式版需防止已確認薪資被隨意重算覆蓋。",
            "文件管理、員工自助、Manager approval、多層審批和通知系統仍在 roadmap。",
        ],
    )

    add_section(doc, "6. 下一步 Roadmap")
    add_table(
        doc,
        ["階段", "目標", "主要工作"],
        [
            ["Phase 1：MVP 穩定化", "把目前可示範功能變得更穩定、更易用。", "補 UI 一致性、mobile tables、權限測試、設定項管理、操作錯誤處理。"],
            ["Phase 2：HR Core", "完善員工主檔和敏感資料管理。", "員工完整資料、緊急聯絡人、銀行/MPF、資料遮罩、資料更改申請。"],
            ["Phase 3：Leave + Attendance", "建立更完整假期及出勤流程。", "年假政策、病假文件、公眾假期日曆、Manager approval、Team calendar、打卡/補鐘。"],
            ["Phase 4：Payroll / MPF", "把薪資流程推向正式可用。", "Payroll run lock、approval flow、正式 payslip PDF、MPF 報表、IR56B 正式化。"],
            ["Phase 5：Onboarding / Offboarding", "支援入職、離職與文件流程。", "入職 checklist、合約模板、文件收集、離職 final pay、資產歸還。"],
            ["Phase 6：Production Ready", "滿足正式部署、安全和維運要求。", "PostgreSQL、migration、backup、2FA、rate limit、error monitoring、automated tests。"],
        ],
        widths=[1.55, 2.2, 3.05],
    )

    add_subsection(doc, "建議近期優先次序")
    add_numbered(
        doc,
        [
            "補齊 UI mobile card pattern：員工、請假、薪資項目、Audit、報表頁。",
            "做權限測試：Employee 只能看自己的資料，HR/Admin 才能改公司設定和薪資資料。",
            "Payroll lock / unlock：避免已確認薪資被重算覆蓋。",
            "IR56B 欄位正式化：對照香港 IRD 欄位與公司薪資項目。",
            "敏感資料遮罩與加密：HKID、銀行戶口、薪資、地址、電話。",
            "Employee Self-Service：員工首頁、個人資料、請假紀錄、糧單下載。",
        ],
    )

    add_section(doc, "7. 客戶溝通重點")
    add_bullets(
        doc,
        [
            "這不是單一 Excel 工具，而是一套可逐步擴展的 HR workflow platform。",
            "目前版本已可展示員工、請假、薪資、MPF、報表、Audit Trail 的完整 MVP 流程。",
            "系統採用可配置 master data，部門、職位、銀行、假期類型和薪資項目不需要寫死在程式中。",
            "香港本地化是產品方向，包括繁體中文、MPF、IR56B、香港假期及中小企 HR 操作習慣。",
            "正式上線前需要進一步處理安全、法規覆核、薪資鎖定、資料保留和備份。",
            "最適合的合作方式是先用 1 至 2 間小型公司做 pilot，收集真實 HR 流程後再產品化。",
        ],
    )

    add_section(doc, "8. 建議客戶 Pilot 範圍")
    add_table(
        doc,
        ["Pilot 項目", "建議範圍"],
        [
            ["公司規模", "5 至 30 名員工"],
            ["使用角色", "1 至 2 名 HR/Admin，少量員工測試 Self-Service"],
            ["測試周期", "2 至 4 星期"],
            ["測試資料", "匿名或測試員工資料，避免一開始放入完整真實敏感資料"],
            ["驗證重點", "員工資料建立、請假申請、薪資計算、MPF、報表、Audit Trail"],
            ["成功標準", "HR 能用系統完成核心流程，並確認哪些公司政策需要配置化"],
        ],
        widths=[1.75, 5.05],
    )

    add_section(doc, "9. 免責與合規說明")
    add_para(doc, "本文件描述的是目前 MVP 系統能力及產品 roadmap，不構成法律、稅務或會計意見。涉及香港《僱傭條例》、MPF、IR56B、個人資料私隱及薪資計算的正式使用，應由香港 HR 專業人士、會計師或法律顧問覆核後再上線。")
    add_para(doc, "正式商業版本應加入更完整的權限控制、資料加密、備份、資料保留政策、操作審批和審計報表。")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Hong Kong SME HR System MVP | Client Presentation Document")
    set_run_font(run, size=9, color=RGBColor(100, 116, 139))
    return doc


if __name__ == "__main__":
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    build_document().save(OUTPUT)
    print(OUTPUT)
